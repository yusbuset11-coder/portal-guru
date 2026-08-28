from datetime import datetime
from io import BytesIO
import docx
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
import google.generativeai as genai
import gspread
from google.oauth2.service_account import Credentials
import pandas as pd
import streamlit as st

# --- KONFIGURASI HALAMAN ---
st.set_page_config(
    page_title="SAKTI - Sistem Asesmen & Kompetensi Terintegrasi",
    page_icon="🎯",
    layout="wide",
)

# --- CEK SESI LOGIN DARI PORTAL UTAMA (app.py) ---
if "logged_in" not in st.session_state or not st.session_state.logged_in:
  st.warning(
      "⚠️ Anda belum login. Silakan kembali ke Halaman Utama (`app.py`) untuk"
      " masuk ke portal terlebih dahulu."
  )
  st.stop()


# --- KONEKSI GOOGLE SHEETS & GEMINI AI ---
@st.cache_resource
def init_connections():
  try:
    scope = [
        "https://www.googleapis.com/auth/spreadsheets",
        "https://www.googleapis.com/auth/drive",
    ]
    creds_dict = dict(st.secrets["gcp_service_account"])
    creds = Credentials.from_service_account_info(creds_dict, scopes=scope)
    client = gspread.authorize(creds)

    genai.configure(api_key=st.secrets["GEMINI_API_KEY"])
    return client
  except Exception as e:
    st.error(f"Gagal terhubung ke sistem Google/Gemini: {e}")
    return None


gc = init_connections()

# Ambil ID Spreadsheet khusus guru dari session state portal utama
user_spreadsheet_id = st.session_state.get("spreadsheet_id", "")


# --- FUNGSI BANTUAN FORMAT TABEL WORD PROFESIONAL ---
def set_cell_background(cell, fill_color):
  tcPr = cell._tc.get_or_add_tcPr()
  shd = OxmlElement("w:shd")
  shd.set(qn("w:val"), "clear")
  shd.set(qn("w:color"), "auto")
  shd.set(qn("w:fill"), fill_color)
  tcPr.append(shd)


def set_cell_margins(cell, top=80, bottom=80, left=120, right=120):
  tcPr = cell._tc.get_or_add_tcPr()
  tcMar = OxmlElement("w:tcMar")
  for m, val in [
      ("top", top),
      ("bottom", bottom),
      ("left", left),
      ("right", right),
  ]:
    node = OxmlElement(f"w:{m}")
    node.set(qn("w:w"), str(val))
    node.set(qn("w:type"), "dxa")
    tcMar.append(node)
  tcPr.append(tcMar)


# --- FUNGSI PEMBUAT WORD PROFESIONAL ---
def generate_professional_word_document(
    mapel, materi, kelas, jenjang, fase, jenis_asesmen, sub_asesmen, jumlah_soal, jenis_soal, content_text
):
  doc = docx.Document()

  for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

  # Judul Dokumen Utama
  p_title = doc.add_paragraph()
  p_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
  run_title = p_title.add_run(f"INSTRUMEN {jenis_asesmen.upper()}")
  run_title.bold = True
  run_title.font.name = "Cambria"
  run_title.font.size = Pt(14)
  run_title.font.color.rgb = RGBColor(27, 54, 93)
  p_title.paragraph_format.space_after = Pt(8)

  # TABEL 1: Informasi / Metadata Instrumen (Jarak Rapat)
  table_meta = doc.add_table(rows=5, cols=2)
  table_meta.alignment = WD_TABLE_ALIGNMENT.CENTER

  metadata = [
      ("Mata Pelajaran", mapel),
      ("Materi/Topik", materi),
      ("Jenjang/Kelas", f"{jenjang} / {fase} ({kelas})"),
      (
          "Tingkat Kesulitan",
          "Sedang (Menguji Kemampuan Analisis, Evaluasi, dan Kontekstualisasi)",
      ),
      ("Jenis Asesmen", jenis_asesmen),  # Label diganti dari Bentuk Soal
  ]

  for i, (key, val) in enumerate(metadata):
    row = table_meta.rows[i]
    cell_key, cell_val = row.cells[0], row.cells[1]
    cell_key.width = Inches(2.2)
    cell_val.width = Inches(4.3)

    p_k = cell_key.paragraphs[0]
    p_k.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_k.paragraph_format.space_before = Pt(2)
    p_k.paragraph_format.space_after = Pt(2)
    p_k.paragraph_format.line_spacing = 1.15
    r_k = p_k.add_run(key)
    r_k.bold = True
    r_k.font.name = "Cambria"
    r_k.font.size = Pt(10.5)

    p_v = cell_val.paragraphs[0]
    p_v.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_v.paragraph_format.space_before = Pt(2)
    p_v.paragraph_format.space_after = Pt(2)
    p_v.paragraph_format.line_spacing = 1.15
    r_v = p_v.add_run(str(val))
    r_v.font.name = "Cambria"
    r_v.font.size = Pt(10.5)

    set_cell_margins(cell_key, 80, 80, 120, 120)
    set_cell_margins(cell_val, 80, 80, 120, 120)
    set_cell_background(cell_key, "F2F4F8")

  doc.add_paragraph().paragraph_format.space_after = Pt(6)

  # TABEL 2: Spesifikasi Asesmen (Tabel Pengganti Bullet Points)
  p_sub = doc.add_paragraph()
  run_sub = p_sub.add_run("SPESIFIKASI ASESMEN")
  run_sub.bold = True
  run_sub.font.name = "Cambria"
  run_sub.font.size = Pt(12)
  run_sub.font.color.rgb = RGBColor(27, 54, 93)
  p_sub.paragraph_format.space_after = Pt(4)

  specs = [
      ("Mata Pelajaran", mapel),
      ("Fase / Kelas", f"{fase} ({kelas})"),
      ("Materi/Topik", materi),
      ("Jumlah Soal", f"{jumlah_soal} Butir"),
      ("Bentuk Soal", jenis_soal),
      ("Sub Jenis Asesmen", sub_asesmen),
  ]

  table_spec = doc.add_table(rows=len(specs), cols=2)
  table_spec.alignment = WD_TABLE_ALIGNMENT.CENTER
  for i, (k, v) in enumerate(specs):
    row = table_spec.rows[i]
    c0, c1 = row.cells[0], row.cells[1]
    c0.width = Inches(2.2)
    c1.width = Inches(4.3)
    c0.text = k
    c1.text = v
    for cell in (c0, c1):
      for p in cell.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.15
        for run in p.runs:
          run.font.name = "Cambria"
          run.font.size = Pt(10.5)
    set_cell_margins(c0, 80, 80, 120, 120)
    set_cell_margins(c1, 80, 80, 120, 120)
    set_cell_background(c0, "F2F4F8")

  doc.add_paragraph().paragraph_format.space_after = Pt(6)

  # BUTIR SOAL TITLE
  p_bs_title = doc.add_paragraph()
  r_bst = p_bs_title.add_run("BUTIR SOAL & PEMBAHASAN")
  r_bst.bold = True
  r_bst.font.name = "Cambria"
  r_bst.font.size = Pt(12)
  r_bst.font.color.rgb = RGBColor(27, 54, 93)
  p_bs_title.paragraph_format.space_after = Pt(6)

  # PARSING KONTEN SOAL DENGAN FORMAT RATA KIRI-KANAN (JUSTIFY)
  for line in content_text.split("\n"):
    line_str = line.strip()
    if not line_str:
      continue

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # Otomatis Justify
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15

    if (
        line_str.startswith("###")
        or line_str.startswith("##")
        or line_str.startswith("#")
    ):
      run = p.add_run(line_str.replace("#", "").strip())
      run.font.name = "Cambria"
      run.font.size = Pt(11.5)
      run.font.bold = True
      run.font.color.rgb = RGBColor(27, 54, 93)
      p.paragraph_format.space_before = Pt(8)
    elif (
        line_str.startswith("SOAL")
        or line_str.startswith("KUNCI")
        or line_str.startswith("Pertanyaan")
        or line_str.startswith("Pembahasan")
        or line_str.startswith("Soal")
    ):
      run = p.add_run(line_str)
      run.font.name = "Cambria"
      run.font.size = Pt(10.5)
      run.font.bold = True
      run.font.color.rgb = RGBColor(15, 23, 42)
      p.paragraph_format.space_before = Pt(4)
    else:
      run = p.add_run(line_str)
      run.font.name = "Cambria"
      run.font.size = Pt(10.5)
      run.font.color.rgb = RGBColor(51, 65, 85)

  bio = BytesIO()
  doc.save(bio)
  bio.seek(0)
  return bio


# --- HEADER APLIKASI ---
st.markdown(
    """
    <div style="background: linear-gradient(135deg, #b45309, #d97706); padding: 20px; border-radius: 15px; color: white; margin-bottom: 25px;">
        <h1 style="margin: 0; font-size: 26px;">🎯 SAKTI</h1>
        <p style="margin: 5px 0 0 0; font-size: 14px; opacity: 0.9;">Sistem Asesmen & Kompetensi Terintegrasi (Pendekatan Pembelajaran Mendalam)</p>
    </div>
""",
    unsafe_allow_html=True,
)

# --- SIDEBAR NAVIGASI ---
with st.sidebar:
  st.success(f"👤 **{st.session_state.get('guru_nama', 'Guru')}**")
  st.markdown("---")
  st.markdown("### 📌 Menu Navigasi SAKTI")
  menu = st.radio(
      "Pilih Menu:",
      [
          "🏠 Beranda SAKTI",
          "✨ Generator Asesmen AI",
          "📊 Bank Soal & Asesmen Tersimpan",
          "📈 Rekap Nilai Siswa",
      ],
  )

# --- MENU 1: BERANDA SAKTI ---
if menu == "🏠 Beranda SAKTI":
  st.subheader(
      f"Selamat Datang, **{st.session_state.get('guru_nama', 'Guru')}** di Modul"
      " SAKTI"
  )
  st.write(
      "Gunakan kecerdasan buatan untuk merancang soal asesmen formatif,"
      " sumatif, kisi-kisi, serta rubrik penilaian secara cepat dan akurat."
  )

  with st.container(border=True):
    st.markdown("### **✨ Fitur Unggulan SAKTI**")
    st.markdown(
        "* **Generator Soal Otomatis:** Buat soal Pilihan Ganda dan Essay"
        " berdasarkan Capaian Pembelajaran (CP) atau materi spesifik dengan"
        " Pendekatan Pembelajaran Mendalam."
    )
    st.markdown(
        "* **Kunci Jawaban & Pembahasan:** Dilengkapi opsi pembahasan"
        " mendalam untuk setiap butir soal."
    )
    st.markdown(
        "* **Penyimpanan Cloud & Word Profesional:** Simpan hasil asesmen ke"
        " Google Spreadsheet dan unduh dokumen Word siap cetak."
    )
    st.markdown(
        "* **Rekap Nilai Siswa:** Kelola dan sinkronkan rekap nilai siswa"
        " langsung terhubung ke database kelas masing-masing."
    )

# --- MENU 2: GENERATOR ASESMEN AI ---
elif menu == "✨ Generator Asesmen AI":
  st.subheader("Parameter Pembuatan Soal & Asesmen (Pendekatan PM)")

  col1, col2 = st.columns(2)
  with col1:
    mapel = st.text_input(
        "Mata Pelajaran", placeholder="Contoh: Bahasa Indonesia"
    )
  with col2:
    materi = st.text_input(
        "Materi / Topik", placeholder="Contoh: Mengidentifikasi Makna Kata"
    )

  col3, col4, col5 = st.columns(3)
  with col3:
    jenjang = st.selectbox(
        "Jenjang",
        ["-- Pilih Jenjang --", "SD", "SMP", "SMA", "SMK"],
        key="gen_jenjang",
    )

  if jenjang == "SD":
    fase_options = [
        "Fase A (Kelas 1-2)",
        "Fase B (Kelas 3-4)",
        "Fase C (Kelas 5-6)",
    ]
    kelas_options = [
        "Kelas 1",
        "Kelas 2",
        "Kelas 3",
        "Kelas 4",
        "Kelas 5",
        "Kelas 6",
    ]
  elif jenjang == "SMP":
    fase_options = ["Fase D (Kelas 7-9)"]
    kelas_options = ["Kelas 7", "Kelas 8", "Kelas 9"]
  elif jenjang in ["SMA", "SMK"]:
    fase_options = ["Fase E (Kelas 10)", "Fase F (Kelas 11-12)"]
    kelas_options = ["Kelas 10", "Kelas 11", "Kelas 12"]
  else:
    fase_options = ["-- Pilih Jenjang Dulu --"]
    kelas_options = ["-- Pilih Jenjang Dulu --"]

  with col4:
    fase = st.selectbox("Fase", fase_options, key="gen_fase")
  with col5:
    kelas = st.selectbox("Kelas", kelas_options, key="gen_kelas")

  col6, col7 = st.columns(2)
  with col6:
    jenis_asesmen = st.selectbox(
        "Jenis Asesmen (Pendekatan PM)",
        ["Asesmen Formatif", "Asesmen Sumatif"],
        key="gen_jenis",
    )
  with col7:
    sub_asesmen = st.selectbox(
        "Bentuk / Sub Jenis Asesmen",
        ["Tulis", "Lisan", "Praktik", "Proyek", "Produk", "Portofolio"],
    )

  col8, col9, col10 = st.columns(3)
  with col8:
    jumlah_soal = st.number_input(
        "Jumlah Butir Soal", min_value=1, max_value=20, value=5
    )
  with col9:
    jenis_soal = st.selectbox(
        "Jenis Soal",
        [
            "Pilihan Ganda",
            "Uraian (Esai)",
            "Jawaban Singkat",
            "Benar - Salah",
            "Menjodohkan",
        ],
    )
  with col10:
    kesulitan = st.selectbox(
        "Tingkat Kesulitan", ["Mudah", "Sedang", "Sulit"], index=1
    )

  if jenjang == "SD":
    aturan_opsi = "3 opsi (A sampai C)"
  elif jenjang == "SMP":
    aturan_opsi = "4 opsi (A sampai D)"
  else:
    aturan_opsi = "5 opsi (A sampai E)"

  if st.button("✨ Buat Instrumen Asesmen dengan Gemini AI 🚀", use_container_width=True):
    if jenjang == "-- Pilih Jenjang --" or not mapel or not materi:
      st.warning(
          "Mohon lengkapi Mata Pelajaran, Materi, dan pilih Jenjang terlebih"
          " dahulu!"
      )
    else:
      with st.spinner(
          "⏳ Sedang merancang instrumen asesmen mendalam dengan Gemini AI..."
      ):
        try:
          prompt = f"""Bertindaklah sebagai pakar kurikulum dan penyusun instrumen asesmen profesional. Buatkan {jumlah_soal} butir soal dengan bentuk **{jenis_soal}**, dalam bentuk asesmen {sub_asesmen} ({jenis_asesmen}) untuk Mata Pelajaran: {mapel}, Materi/Topik: {materi}, Jenjang: {jenjang} ({fase}, {kelas}), dengan tingkat kesulitan {kesulitan}. 
                    
                    Ketentuan Khusus:
                    1. Gunakan pendekatan Pembelajaran Mendalam (Deep Learning) yang merangsang berpikir kritis dan kontekstual.
                    2. Jika jenis soal adalah Pilihan Ganda, gunakan {aturan_opsi}.
                    3. Berikan kunci jawaban yang jelas serta pembahasan mendalam untuk setiap soal."""

          model = genai.GenerativeModel("gemini-1.5-flash")
          response = model.generate_content(prompt)

          st.session_state.generated_soal = response.text
          st.session_state.judul_soal = f"Asesmen_{mapel}_{materi}"
          st.session_state.meta_mapel = mapel
          st.session_state.meta_materi = materi
          st.session_state.meta_kelas = kelas
          st.session_state.meta_jenjang = jenjang
          st.session_state.meta_fase = fase
          st.session_state.meta_jenis = jenis_asesmen
          st.session_state.meta_sub_asesmen = sub_asesmen
          st.session_state.meta_jumlah_soal = jumlah_soal
          st.session_state.meta_jenis_soal = jenis_soal

          st.success("✅ Instrumen Asesmen Berhasil Dibuat!")

        except Exception as e:
          st.error(f"Terjadi kesalahan saat memanggil AI: {e}")

  if "generated_soal" in st.session_state and st.session_state.generated_soal:
    st.markdown("### 📋 Hasil Instrumen Asesmen dari AI:")
    with st.container(border=True):
      st.markdown(st.session_state.generated_soal)

    col_dl1, col_dl2 = st.columns(2)

    with col_dl1:
      docx_file = generate_professional_word_document(
          mapel=st.session_state.get("meta_mapel", mapel),
          materi=st.session_state.get("meta_materi", materi),
          kelas=st.session_state.get("meta_kelas", kelas),
          jenjang=st.session_state.get("meta_jenjang", jenjang),
          fase=st.session_state.get("meta_fase", fase),
          jenis_asesmen=st.session_state.get("meta_jenis", jenis_asesmen),
          sub_asesmen=st.session_state.get("meta_sub_asesmen", sub_asesmen),
          jumlah_soal=st.session_state.get("meta_jumlah_soal", jumlah_soal),
          jenis_soal=st.session_state.get("meta_jenis_soal", jenis_soal),
          content_text=st.session_state.generated_soal,
      )

      st.download_button(
          label="📥 Download Soal Format Word (Siap Cetak)",
          data=docx_file,
          file_name=f"{st.session_state.judul_soal}.docx",
          mime=(
              "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
          ),
          use_container_width=True,
      )

    with col_dl2:
      if st.button("💾 Simpan ke Bank Soal Spreadsheet", use_container_width=True):
        if user_spreadsheet_id:
          with st.spinner("Menyimpan ke Google Sheets..."):
            try:
              ss_user = gc.open_by_key(user_spreadsheet_id)
              try:
                ws_bank = ss_user.worksheet("Bank Soal SAKTI")
              except Exception:
                ws_bank = ss_user.add_worksheet(
                    title="Bank Soal SAKTI", rows="1000", cols="5"
                )
                ws_bank.append_row([
                    "Tanggal",
                    "Mata_Pelajaran",
                    "Materi",
                    "Jenis_Asesmen",
                    "Konten_Soal",
                ])

              ws_bank.append_row([
                  str(datetime.today().strftime("%Y-%m-%d")),
                  str(st.session_state.get("meta_mapel", "")),
                  str(st.session_state.get("meta_materi", "")),
                  str(st.session_state.get("meta_jenis", "")),
                  str(st.session_state.generated_soal),
              ])
              st.success(
                  "🎉 Soal berhasil disimpan ke tab 'Bank Soal SAKTI' di"
                  " spreadsheet Anda!"
              )
            except Exception as e:
              st.error(f"Gagal menyimpan ke spreadsheet: {e}")

# --- MENU 3: BANK SOAL & ASESMEN TERSIMPAN ---
elif menu == "📊 Bank Soal & Asesmen Tersimpan":
  st.subheader("📊 **Bank Soal & Asesmen Tersimpan**")
  st.write("Daftar arsip naskah asesmen yang pernah Anda buat dan simpan.")

  if user_spreadsheet_id:
    try:
      ss_user = gc.open_by_key(user_spreadsheet_id)
      ws_bank = ss_user.worksheet("Bank Soal SAKTI")
      data_bank = ws_bank.get_all_records()

      if not data_bank:
        st.info("Belum ada arsip soal yang tersimpan di database.")
      else:
        df_bank = pd.DataFrame(data_bank)
        with st.container(border=True):
          st.dataframe(df_bank, use_container_width=True)
    except Exception:
      st.info(
          "Tab 'Bank Soal SAKTI' belum tersedia atau belum ada data di"
          " spreadsheet Anda."
      )

# --- MENU 4: REKAP NILAI SISWA ---
elif menu == "📈 Rekap Nilai Siswa":
  st.subheader("Simpan & Rekap Nilai Hasil Asesmen")

  if not user_spreadsheet_id:
    st.error(
        "⚠️ Spreadsheet ID untuk akun Anda tidak ditemukan di Master Registry."
        " Hubungi administrator."
    )
    st.stop()


  @st.cache_data(ttl=5)
  def fetch_master_data(sheet_id):
    try:
      ss = gc.open_by_key(sheet_id)
      ws = ss.worksheet("Data Kelas-Siswa")
      records = ws.get_all_records()
      return records
    except Exception as e:
      return []


  master_data = fetch_master_data(user_spreadsheet_id)

  daftar_sekolah = []
  for r in master_data:
    sekolah_val = str(r.get("Sekolah", "")).strip()
    if sekolah_val and sekolah_val not in daftar_sekolah:
      daftar_sekolah.append(sekolah_val)

  if not daftar_sekolah:
    daftar_sekolah = ["SMK Negeri 2 Bangkalan"]

  if "val_mapel" not in st.session_state:
    st.session_state.val_mapel = ""
  if "val_sekolah" not in st.session_state:
    st.session_state.val_sekolah = daftar_sekolah[0]
  if "val_kelas" not in st.session_state:
    st.session_state.val_kelas = ""
  if "val_jenis" not in st.session_state:
    st.session_state.val_jenis = "Asesmen Formatif"
  if "val_materi" not in st.session_state:
    st.session_state.val_materi = ""

  col_r1, col_r2 = st.columns(2)
  with col_r1:
    tanggal_input = st.date_input("Pilih Tanggal", value=datetime.now())
    r_mapel = st.text_input(
        "Mata Pelajaran",
        value=st.session_state.val_mapel,
        placeholder="Contoh: Bahasa Indonesia",
    )
    r_sekolah = st.selectbox(
        "Pilih Sekolah", daftar_sekolah, key="input_sekolah"
    )

    kelas_filtered = [
        r
        for r in master_data
        if str(r.get("Sekolah", "")).strip() == r_sekolah
    ]
    daftar_kelas = []
    for r in kelas_filtered:
      kelas_val = str(r.get("Kelas", "")).strip()
      if kelas_val and kelas_val not in daftar_kelas:
        daftar_kelas.append(kelas_val)

    if not daftar_kelas:
      daftar_kelas = ["X TKR-1", "X DKV-1"]

    r_kelas = st.selectbox(
        "Pilih Kelas", sorted(daftar_kelas), key="input_kelas"
    )
    r_jenis = st.selectbox(
        "Jenis Asesmen",
        ["Asesmen Formatif", "Asesmen Sumatif"],
        key="input_jenis",
    )

  with col_r2:
    r_materi = st.text_input(
        "Materi / Topik",
        value=st.session_state.val_materi,
        placeholder="Contoh: Teks LHO",
    )

    siswa_filtered = [
        r
        for r in master_data
        if str(r.get("Sekolah", "")).strip() == r_sekolah
        and str(r.get("Kelas", "")).strip() == r_kelas
    ]

    mapping_absen_nama = {}
    for r in siswa_filtered:
      try:
        absen_raw = r.get("No_Absen", r.get("No Absen", 1))
        nama_raw = r.get("Nama_Siswa", r.get("Nama Siswa", r.get("Nama", "")))

        absen = int(absen_raw)
        nama = str(nama_raw).strip()
        if nama:
          mapping_absen_nama[absen] = nama
      except:
        continue

    if not mapping_absen_nama:
      mapping_absen_nama = {1: "Siswa 1", 2: "Siswa 2"}

    list_absen = sorted(list(mapping_absen_nama.keys()))
    r_no_absen = st.selectbox("Pilih No. Absen Siswa", list_absen)

    r_nama_siswa = mapping_absen_nama.get(r_no_absen, "")
    st.text_input(
        "Nama Siswa (Otomatis dari Spreadsheet)",
        value=r_nama_siswa,
        disabled=True,
    )

    r_nilai = st.number_input(
        "Nilai Siswa", min_value=0, max_value=100, value=80
    )

  if st.button("💾 Simpan Nilai ke Google Sheets", use_container_width=True):
    if not r_mapel or not r_nama_siswa:
      st.warning(
          "Mohon lengkapi Mata Pelajaran dan pastikan Nama Siswa terpilih."
      )
    else:
      try:
        ss = gc.open_by_key(user_spreadsheet_id)

        correct_header = [
            "Tanggal",
            "Nama Guru",
            "Sekolah",
            "Mata Pelajaran",
            "Jenjang",
            "Kelas",
            "Jenis Asesmen",
            "Sub Jenis Asesmen",
            "Materi",
            "No Absen",
            "Nama Siswa",
            "Nilai",
        ]

        try:
          sheet = ss.worksheet("Rekap_Nilai")
          sheet.update("A1:L1", [correct_header])
        except:
          sheet = ss.add_worksheet(title="Rekap_Nilai", rows=100, cols=12)
          sheet.append_row(correct_header)

        tanggal_str = tanggal_input.strftime("%Y-%m-%d")

        sheet.append_row([
            tanggal_str,
            st.session_state.get("guru_nama", ""),
            r_sekolah,
            r_mapel,
            "",
            r_kelas,
            r_jenis,
            "",
            r_materi,
            r_no_absen,
            r_nama_siswa,
            r_nilai,
        ])

        st.session_state.val_mapel = r_mapel
        st.session_state.val_sekolah = r_sekolah
        st.session_state.val_kelas = r_kelas
        st.session_state.val_jenis = r_jenis
        st.session_state.val_materi = r_materi

        st.success(
            f"🎉 Nilai untuk **{r_nama_siswa}** (Absen: {r_no_absen}) berhasil"
            " disimpan ke spreadsheet pribadi Anda!"
        )
        st.balloons()
      except Exception as e:
        st.error(f"Gagal menyimpan ke database: {e}")

  # --- BAGIAN MENU UNDUH / EXPORT REKAP NILAI ---
  st.markdown("---")
  st.subheader("📥 Unduh Rekap Nilai Siswa")
  st.markdown(
      "Pilih filter sekolah, kelas, dan mata pelajaran untuk mengunduh rekap"
      " nilai ke dalam format Excel."
  )

  try:
    ss_rekap = gc.open_by_key(user_spreadsheet_id)
    ws_rekap = ss_rekap.worksheet("Rekap_Nilai")
    all_rekap_records = ws_rekap.get_all_records()
  except:
    all_rekap_records = []

  if all_rekap_records:
    df_rekap = pd.DataFrame(all_rekap_records)

    list_dl_sekolah = (
        df_rekap["Sekolah"].unique().tolist()
        if "Sekolah" in df_rekap.columns
        else daftar_sekolah
    )
    dl_sekolah = st.selectbox(
        "Filter Sekolah untuk Unduh", list_dl_sekolah, key="dl_sek"
    )

    df_filtered_sek = (
        df_rekap[df_rekap["Sekolah"] == dl_sekolah]
        if "Sekolah" in df_rekap.columns
        else df_rekap
    )
    list_dl_kelas = (
        df_filtered_sek["Kelas"].unique().tolist()
        if "Kelas" in df_filtered_sek.columns
        else []
    )
    dl_kelas = st.selectbox(
        "Filter Kelas untuk Unduh",
        list_dl_kelas if list_dl_kelas else ["Semua Kelas"],
        key="dl_kls",
    )

    df_filtered_kls = (
        df_filtered_sek[df_filtered_sek["Kelas"] == dl_kelas]
        if dl_kelas != "Semua Kelas"
        else df_filtered_sek
    )
    list_dl_mapel = (
        df_filtered_kls["Mata Pelajaran"].unique().tolist()
        if "Mata Pelajaran" in df_filtered_kls.columns
        else []
    )
    dl_mapel = st.selectbox(
        "Filter Mata Pelajaran untuk Unduh",
        list_dl_mapel if list_dl_mapel else ["Semua Mapel"],
        key="dl_mpl",
    )

    df_final_dl = df_filtered_kls
    if dl_mapel != "Semua Mapel":
      df_final_dl = df_final_dl[df_final_dl["Mata Pelajaran"] == dl_mapel]

    st.write(
        f"📊 Menemukan **{len(df_final_dl)}** data nilai sesuai filter yang"
        " dipilih."
    )
    if not df_final_dl.empty:
      st.dataframe(df_final_dl, use_container_width=True)

      output_excel = BytesIO()
      with pd.ExcelWriter(output_excel, engine="openpyxl") as writer:
        df_final_dl.to_excel(writer, index=False, sheet_name="Rekap_Nilai")
      output_excel.seek(0)

      st.download_button(
          label="📥 Download Data Rekap Terpilih (.xlsx)",
          data=output_excel,
          file_name=(
              f"Rek_Nilai_{dl_sekolah}_{dl_kelas}_{dl_mapel}.xlsx".replace(
                  " ", "_"
              )
          ),
          mime=(
              "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
          ),
          use_container_width=True,
      )
    else:
      st.info(
          "Tidak ada data nilai yang cocok dengan kombinasi filter tersebut."
      )
  else:
    st.info(
        "Belum ada data rekap nilai yang tersimpan di spreadsheet Anda. Silakan"
        " simpan beberapa nilai terlebih dahulu."
    )