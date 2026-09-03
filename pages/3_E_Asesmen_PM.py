from datetime import datetime
from io import BytesIO
import docx
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
import google.generativeai as genai
import gspread
from google.oauth2.service_account import Credentials
import pandas as pd
import streamlit as st
from styles import apply_global_styles

# --- KONFIGURASI HALAMAN ---
st.set_page_config(
    page_title="E Asesmen Pembelajaran Mendalam - Sistem Asesmen & Kompetensi",
    page_icon="🎯",
    layout="wide",
)
apply_global_styles()

st.markdown("""
    <style>
        .block-container {
            padding-top: 1.5rem;
            padding-bottom: 7rem;
        }
    </style>
""", unsafe_allow_html=True)

# --- CEK SESI LOGIN DARI PORTAL UTAMA (app.py) ---
if "logged_in" not in st.session_state or not st.session_state.logged_in:
    st.warning("⚠️ Anda belum login. Silakan kembali ke Halaman Utama (`app.py`) untuk masuk ke portal terlebih dahulu.")
    st.stop()

# --- KONEKSI GOOGLE SHEETS & GEMINI AI ---
@st.cache_resource
def init_connections():
    try:
        scope = [
            "https://www.googleapis.com/auth/spreadsheets",
            "https://www.googleapis.com/auth/drive"
        ]
        creds_dict = dict(st.secrets["gcp_service_account"])
        creds = Credentials.from_service_account_info(creds_dict, scopes=scope)
        client = gspread.authorize(creds)
        
        genai.configure(api_key=st.secrets["GEMINI_API_KEY"])
        return client
    except Exception as e:
        st.error(f"Gagal terhubung ke sistem Google/Gemini: {e}")
        return None

gc = init_connections()
user_spreadsheet_id = st.session_state.get("spreadsheet_id", "")

# --- FUNGSI BANTUAN XML WORD (PROFESIONAL & EYE-CATCHING) ---
def set_cell_background(cell, fill_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_color)
    tcPr.append(shd)

def set_cell_margins(cell, top=120, bottom=120, left=180, right=180):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_callout_borders(cell, border_color="1B365D"):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), '24') # 3pt width
    left.set(qn('w:space'), '0')
    left.set(qn('w:color'), border_color)
    tcBorders.append(left)
    
    for b_name in ['top', 'bottom', 'right']:
        b = OxmlElement(f'w:{b_name}')
        b.set(qn('w:val'), 'none')
        tcBorders.append(b)
        
    tcPr.append(tcBorders)

# --- FUNGSI PEMBUAT WORD PROFESIONAL & EYE-CATCHING ---
def generate_professional_word_document(
    mapel, materi, kelas, jenjang, fase, jenis_asesmen, sub_asesmen, jumlah_soal, jenis_soal, content_text
):
    doc = docx.Document()

    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run(f"INSTRUMEN {jenis_asesmen.upper()}")
    run_title.bold = True
    run_title.font.name = "Cambria"
    run_title.font.size = Pt(16)
    run_title.font.color.rgb = RGBColor(27, 54, 93)
    p_title.paragraph_format.space_after = Pt(2)

    p_sub_title = doc.add_paragraph()
    p_sub_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub_title = p_sub_title.add_run(f"MATA PELAJARAN: {mapel.upper()}")
    run_sub_title.font.name = "Cambria"
    run_sub_title.font.size = Pt(12)
    run_sub_title.font.bold = True
    run_sub_title.font.color.rgb = RGBColor(74, 85, 104)
    p_sub_title.paragraph_format.space_after = Pt(14)

    table_meta = doc.add_table(rows=5, cols=2)
    table_meta.alignment = WD_TABLE_ALIGNMENT.CENTER

    metadata = [
        ("Mata Pelajaran", mapel),
        ("Materi / Topik", materi),
        ("Jenjang / Kelas", f"{jenjang} / {fase} ({kelas})"),
        ("Tingkat Kesulitan", "Sedang (Analisis, Evaluasi, & Kontekstualisasi)"),
        ("Jenis & Bentuk Asesmen", f"{jenis_asesmen} ({sub_asesmen}) - {jenis_soal}"),
    ]

    for i, (key, val) in enumerate(metadata):
        row = table_meta.rows[i]
        cell_key, cell_val = row.cells[0], row.cells[1]
        cell_key.width = Inches(2.3)
        cell_val.width = Inches(4.2)

        p_k = cell_key.paragraphs[0]
        r_k = p_k.add_run(key)
        r_k.bold = True
        r_k.font.name = "Cambria"
        r_k.font.size = Pt(10)
        r_k.font.color.rgb = RGBColor(27, 54, 93)

        p_v = cell_val.paragraphs[0]
        r_v = p_v.add_run(str(val))
        r_v.font.name = "Cambria"
        r_v.font.size = Pt(10)
        r_v.font.color.rgb = RGBColor(51, 65, 85)

        set_cell_margins(cell_key, 80, 80, 120, 120)
        set_cell_margins(cell_val, 80, 80, 120, 120)
        set_cell_background(cell_key, "F0F4F8")
        set_cell_background(cell_val, "FFFFFF")

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    p_sec = doc.add_paragraph()
    r_sec = p_sec.add_run("DOKUMEN BUTIR SOAL & PEMBAHASAN")
    r_sec.bold = True
    r_sec.font.name = "Cambria"
    r_sec.font.size = Pt(13)
    r_sec.font.color.rgb = RGBColor(27, 54, 93)
    p_sec.paragraph_format.space_after = Pt(8)

    lines = content_text.split("\n")
    in_callout = False
    callout_buffer = []

    def flush_callout():
        nonlocal callout_buffer
        if callout_buffer:
            tbl = doc.add_table(rows=1, cols=1)
            tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
            c = tbl.rows[0].cells[0]
            c.width = Inches(6.5)
            set_cell_background(c, "F8FAFC")
            set_callout_borders(c, "1B365D")
            set_cell_margins(c, 140, 140, 180, 180)
            
            p_first = True
            for text_line in callout_buffer:
                if p_first:
                    p = c.paragraphs[0]
                    p_first = False
                else:
                    p = c.add_paragraph()
                p.paragraph_format.space_after = Pt(3)
                p.paragraph_format.line_spacing = 1.15
                r = p.add_run(text_line.replace(">", "").strip())
                r.font.name = "Cambria"
                r.font.size = Pt(10)
                r.font.italic = True
                r.font.color.rgb = RGBColor(51, 65, 85)
            callout_buffer = []
            doc.add_paragraph().paragraph_format.space_after = Pt(4)

    for line in lines:
        line_str = line.strip()
        if not line_str:
            if in_callout:
                flush_callout()
                in_callout = False
            continue

        if line_str.startswith(">"):
            in_callout = True
            callout_buffer.append(line_str)
            continue
        else:
            if in_callout:
                flush_callout()
                in_callout = False

        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.15

        if line_str.startswith("SOAL") or line_str.startswith("###") or "---" in line_str:
            if "---" in line_str:
                continue
            p.paragraph_format.space_before = Pt(10)
            run = p.add_run(line_str.replace("#", "").strip())
            run.font.name = "Cambria"
            run.font.size = Pt(11)
            run.font.bold = True
            run.font.color.rgb = RGBColor(27, 54, 93)
        elif line_str.startswith("Kunci Jawaban"):
            p.paragraph_format.space_before = Pt(4)
            run = p.add_run(line_str)
            run.font.name = "Cambria"
            run.font.size = Pt(10.5)
            run.font.bold = True
            run.font.color.rgb = RGBColor(10, 128, 67)
        elif line_str.startswith("Pembahasan"):
            run = p.add_run(line_str)
            run.font.name = "Cambria"
            run.font.size = Pt(10.5)
            run.font.bold = True
            run.font.color.rgb = RGBColor(180, 83, 9)
        elif line_str.startswith("A.") or line_str.startswith("B.") or line_str.startswith("C.") or line_str.startswith("D.") or line_str.startswith("E."):
            p.paragraph_format.left_indent = Inches(0.25)
            run = p.add_run(line_str)
            run.font.name = "Cambria"
            run.font.size = Pt(10.5)
            run.font.color.rgb = RGBColor(51, 65, 85)
        else:
            run = p.add_run(line_str)
            run.font.name = "Cambria"
            run.font.size = Pt(10.5)
            run.font.color.rgb = RGBColor(51, 65, 85)

    if in_callout:
        flush_callout()

    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio

# HEADER UTAMA
guru_nama = st.session_state.get("guru_nama", "")

st.markdown(
    f"""
    <div style="background: linear-gradient(135deg, #1e293b 0%, #111827 100%); padding: 25px 30px; border-radius: 14px; border: 1px solid #334155; margin-bottom: 25px;">
        <h2 style="color: #ffffff; margin: 0 0 10px 0; font-size: 26px;">🎯 Digitalisasi Asesmen Pembelajaran Mendalam</h2>
        <p style="color: #e2e8f0; font-size: 15px; margin: 0 0 8px 0;">Selamat Datang, {guru_nama} di Modul E Asesmen PM</p>
        <p style="color: #94a3b8; font-size: 14px; margin: 0;">Gunakan kecerdasan buatan untuk merancang soal asesmen formatif, sumatif, kisi-kisi, serta rubrik penilaian secara cepat dan akurat.</p>
    </div>
    """,
    unsafe_allow_html=True,
)

# --- SIDEBAR NAVIGASI ---
with st.sidebar:
    st.markdown(
        f"""
        <div class="user-profile-box">
            <span style="font-size: 24px;">👨‍💻</span><br>
            <b style="color: #facc15; font-size: 14px;">{st.session_state.guru_nama}</b><br>
            <span style="color: #94a3b8; font-size: 11px;">Sesi Aktif & Terverifikasi</span>
        </div>
        """,
        unsafe_allow_html=True,
    )

    st.markdown("---")
    menu = st.selectbox(
    "Pilih Menu Asesmen",
    [
        "🏠 Beranda Asesmen",
        "🤖 Generator Asesmen AI",
        "💾 Bank Soal & Asesmen Tersimpan",
        "📊 Input dan Rekap Nilai Siswa",
    ],
)

# --- FUNGSI AMAN BANK SOAL SAKTI ---
@st.cache_data(ttl=10)
def get_clean_bank_soal_dataframe(sheet_id):
    try:
        ss_user = gc.open_by_key(sheet_id)
        try:
            ws = ss_user.worksheet("Bank Soal SAKTI")
        except Exception:
            ws = ss_user.worksheet("Bank Soal")
        rows = ws.get_all_values()
        if not rows or len(rows) <= 1:
            return pd.DataFrame()
        df = pd.DataFrame(rows[1:], columns=rows[0])
        return df
    except Exception as e:
        return pd.DataFrame()

# --- MENU 1: BERANDA ASESMEN ---
if menu == "🏠 Beranda Asesmen":
    st.write("Gunakan kecerdasan buatan untuk merancang soal asesmen formatif, sumatif, kisi-kisi, serta rubrik penilaian secara cepat dan akurat.")

    with st.container(border=True):
        st.markdown("### **✨ Fitur Unggulan Asesmen**")
        st.markdown("* **Generator Soal Otomatis:** Buat soal Pilihan Ganda dan Essay berdasarkan Capaian Pembelajaran (CP) atau materi spesifik dengan Pendekatan Pembelajaran Mendalam.")
        st.markdown("* **Kunci Jawaban & Pembahasan:** Dilengkapi opsi pembahasan mendalam untuk setiap butir soal.")
        st.markdown("* **Penyimpanan Cloud & Word Profesional:** Simpan ringkasan asesmen ke Google Spreadsheet dan unduh dokumen Word siap cetak.")
        st.markdown("* **Rekap Nilai Siswa:** Kelola dan sinkronkan rekap nilai siswa langsung terhubung ke database kelas masing-masing.")

# --- MENU 2: GENERATOR ASESMEN AI ---
elif menu == "✨ Generator Asesmen AI":
    st.subheader("Parameter Pembuatan Soal & Asesmen (Pendekatan PM)")

    col1, col2 = st.columns(2)
    with col1:
        mapel = st.text_input("Mata Pelajaran", placeholder="Contoh: Bahasa Indonesia")
    with col2:
        materi = st.text_input("Materi / Topik", placeholder="Contoh: Mengidentifikasi Makna Kata")

    col3, col4, col5 = st.columns(3)
    with col3:
        jenjang = st.selectbox("Jenjang", ["-- Pilih Jenjang --", "SD", "SMP", "SMA", "SMK"], key="gen_jenjang")

    if jenjang == "SD":
        fase_options = ["Fase A (Kelas 1-2)", "Fase B (Kelas 3-4)", "Fase C (Kelas 5-6)"]
        kelas_options = ["Kelas 1", "Kelas 2", "Kelas 3", "Kelas 4", "Kelas 5", "Kelas 6"]
    elif jenjang == "SMP":
        fase_options = ["Fase D (Kelas 7-9)"]
        kelas_options = ["Kelas 7", "Kelas 8", "Kelas 9"]
    elif jenjang in ["SMA", "SMK"]:
        fase_options = ["Fase E (Kelas 10)", "Fase F (Kelas 11-12)"]
        kelas_options = ["Kelas 10", "Kelas 11", "Kelas 12"]
    else:
        fase_options = ["-- Pilih Jenjang Dulu --"]
        kelas_options = ["-- Pilih Jenjang Dulu --"]

    with col4:
        fase = st.selectbox("Fase", fase_options, key="gen_fase")
    with col5:
        kelas = st.selectbox("Kelas", kelas_options, key="gen_kelas")

    col6, col7 = st.columns(2)
    with col6:
        jenis_asesmen = st.selectbox(
            "Jenis Asesmen (Pendekatan PM)",
            ["Asesmen Formatif", "Asesmen Sumatif"],
            key="gen_jenis",
        )
    with col7:
        if jenis_asesmen == "Asesmen Formatif":
            sub_asesmen_options = ["Formatif Tertulis", "Formatif Tidak Tertulis"]
        else:
            sub_asesmen_options = ["Tulis", "Lisan", "Tugas", "Praktik", "Proyek", "Produk"]

        sub_asesmen = st.selectbox("Bentuk / Sub Jenis Asesmen", sub_asesmen_options)

    col8, col9, col10 = st.columns(3)
    with col8:
        jumlah_soal = st.number_input("Jumlah Butir Soal", min_value=1, max_value=20, value=5)
    with col9:
        jenis_soal = st.selectbox(
            "Jenis Soal",
            ["Pilihan Ganda", "Uraian (Esai)", "Jawaban Singkat", "Benar - Salah", "Menjodohkan"],
        )
    with col10:
        kesulitan = st.selectbox("Tingkat Kesulitan", ["Mudah", "Sedang", "Sulit"], index=1)

    if jenjang == "SD":
        aturan_opsi = "3 opsi (A sampai C)"
    elif jenjang == "SMP":
        aturan_opsi = "4 opsi (A sampai D)"
    else:
        aturan_opsi = "5 opsi (A sampai E)"

    if st.button("✨ Buat Instrumen Asesmen dengan Gemini AI 🚀", use_container_width=True):
        if jenjang == "-- Pilih Jenjang --" or not mapel or not materi:
            st.warning("Mohon lengkapi Mata Pelajaran, Materi, dan pilih Jenjang terlebih dahulu!")
        else:
            with st.spinner("⏳ Sedang merancang instrumen asesmen mendalam dengan Gemini AI..."):
                try:
                    prompt = f"""Bertindaklah sebagai pakar kurikulum dan penyusun instrumen asesmen profesional. Buatkan {jumlah_soal} butir soal dengan bentuk **{jenis_soal}**, dalam bentuk asesmen {sub_asesmen} ({jenis_asesmen}) untuk Mata Pelajaran: {mapel}, Materi/Topik: {materi}, Jenjang: {jenjang} ({fase}, {kelas}), dengan tingkat kesulitan {kesulitan}. 
                    
                    Ketentuan Khusus:
                    1. Gunakan pendekatan Pembelajaran Mendalam (Deep Learning) yang merangsang berpikir kritis dan kontekstual.
                    2. Jika jenis soal adalah Pilihan Ganda, gunakan {aturan_opsi}.
                    3. Berikan kunci jawaban yang jelas serta pembahasan mendalam untuk setiap soal."""

                    model = genai.GenerativeModel("gemini-3.5-flash")
                    response = model.generate_content(prompt)

                    st.session_state.generated_soal = response.text
                    st.session_state.judul_soal = f"Asesmen_{mapel}_{materi}"
                    st.session_state.meta_mapel = mapel
                    st.session_state.meta_materi = materi
                    st.session_state.meta_kelas = kelas
                    st.session_state.meta_jenjang = jenjang
                    st.session_state.meta_fase = fase
                    st.session_state.meta_jenis = jenis_asesmen
                    st.session_state.meta_sub_asesmen = sub_asesmen
                    st.session_state.meta_jumlah_soal = jumlah_soal
                    st.session_state.meta_jenis_soal = jenis_soal

                    st.success("✅ Instrumen Asesmen Berhasil Dibuat!")

                except Exception as e:
                    st.error(f"Terjadi kesalahan saat memanggil AI: {e}")

    if "generated_soal" in st.session_state and st.session_state.generated_soal:
        st.markdown("### 📋 Hasil Instrumen Asesmen dari AI:")
        with st.container(border=True):
            st.markdown(st.session_state.generated_soal)

        col_dl1, col_dl2 = st.columns(2)

        with col_dl1:
            docx_file = generate_professional_word_document(
                mapel=st.session_state.get("meta_mapel", mapel),
                materi=st.session_state.get("meta_materi", materi),
                kelas=st.session_state.get("meta_kelas", kelas),
                jenjang=st.session_state.get("meta_jenjang", jenjang),
                fase=st.session_state.get("meta_fase", fase),
                jenis_asesmen=st.session_state.get("meta_jenis", jenis_asesmen),
                sub_asesmen=st.session_state.get("meta_sub_asesmen", sub_asesmen),
                jumlah_soal=st.session_state.get("meta_jumlah_soal", jumlah_soal),
                jenis_soal=st.session_state.get("meta_jenis_soal", jenis_soal),
                content_text=st.session_state.generated_soal,
            )

            st.download_button(
                label="📥 Download Soal Format Word (Siap Cetak)",
                data=docx_file,
                file_name=f"{st.session_state.judul_soal}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
            )

        with col_dl2:
            if st.button("💾 Simpan ke Bank Soal Spreadsheet", use_container_width=True):
                if user_spreadsheet_id:
                    with st.spinner("Menyimpan ke Google Sheets..."):
                        try:
                            ss_user = gc.open_by_key(user_spreadsheet_id)
                            try:
                                ws_bank = ss_user.worksheet("Bank Soal SAKTI")
                            except Exception:
                                ws_bank = ss_user.add_worksheet(title="Bank Soal SAKTI", rows="1000", cols="10")

                            correct_bank_header = [
                                "Tanggal", "Mata_Pelajaran", "Materi", "Jenjang", "Fase", 
                                "Kelas", "Jenis_Asesmen", "Sub_Jenis_Asesmen", "Jumlah_Soal", 
                                "Jenis_Soal"
                            ]
                            try:
                                ws_bank.update("A1:J1", [correct_bank_header])
                            except:
                                pass

                            ws_bank.append_row([
                                str(datetime.today().strftime("%Y-%m-%d")),
                                str(st.session_state.get("meta_mapel", "")),
                                str(st.session_state.get("meta_materi", "")),
                                str(st.session_state.get("meta_jenjang", "")),
                                str(st.session_state.get("meta_fase", "")),
                                str(st.session_state.get("meta_kelas", "")),
                                str(st.session_state.get("meta_jenis", "")),
                                str(st.session_state.get("meta_sub_asesmen", "")),
                                str(st.session_state.get("meta_jumlah_soal", 5)),
                                str(st.session_state.get("meta_jenis_soal", ""))
                            ])
                            st.success("🎉 Metadata soal berhasil disimpan ke tab 'Bank Soal SAKTI' di spreadsheet Anda!")
                        except Exception as e:
                            st.error(f"Gagal menyimpan ke spreadsheet: {e}")

# --- MENU 3: BANK SOAL & ASESMEN TERSIMPAN ---
elif menu == "📁 Bank Soal & Asesmen Tersimpan":
    st.subheader("📁 **Bank Soal & Asesmen Tersimpan**")
    st.write("Daftar arsip ringkasan asesmen yang pernah Anda buat dan simpan. Klik tombol **Unduh Soal** pada kolom aksi untuk men-generate dan mengunduh dokumen Word siap cetak.")

    if user_spreadsheet_id:
        with st.spinner("Memuat data bank soal..."):
            df_bank = get_clean_bank_soal_dataframe(user_spreadsheet_id)

            if df_bank.empty or len(df_bank) == 0:
                st.info("Belum ada data bank soal yang tersimpan di database.")
            else:
                st.markdown("---")
                h_cols = st.columns([0.5, 1.2, 1.8, 1.8, 1.4, 1.5])
                h_cols[0].markdown("**No.**")
                h_cols[1].markdown("**Tanggal**")
                h_cols[2].markdown("**Mata Pelajaran**")
                h_cols[3].markdown("**Materi**")
                h_cols[4].markdown("**Jenis Asesmen**")
                h_cols[5].markdown("**Aksi**")
                st.markdown("---")

                for idx, row in df_bank.iterrows():
                    nomor = idx + 1
                    tgl = row.get("Tanggal", "")
                    mapel = row.get("Mata_Pelajaran", "Mapel")
                    materi = row.get("Materi", "Materi")
                    jenjang = row.get("Jenjang", "SD")
                    fase = row.get("Fase", "Fase A")
                    kelas = row.get("Kelas", "Kelas 1")
                    jenis_asesmen = row.get("Jenis_Asesmen", "Asesmen Formatif")
                    sub_asesmen = row.get("Sub_Jenis_Asesmen", "Tertulis")
                    try:
                        jumlah_soal = int(row.get("Jumlah_Soal", 5))
                    except:
                        jumlah_soal = 5
                    jenis_soal = row.get("Jenis_Soal", "Pilihan Ganda")

                    row_cols = st.columns([0.5, 1.2, 1.8, 1.8, 1.4, 1.5])
                    row_cols[0].write(str(nomor))
                    row_cols[1].write(tgl)
                    row_cols[2].write(mapel)
                    row_cols[3].write(materi)
                    row_cols[4].write(jenis_asesmen)
                    
                    with row_cols[5]:
                        # Tombol Unduh Soal interaktif untuk setiap baris
                        btn_key = f"dl_bank_{idx}"
                        if st.button("📥 Unduh Soal", key=btn_key, use_container_width=True):
                            with st.spinner("⏳ Menyiapkan dokumen Word dari arsip..."):
                                try:
                                    if jenjang == "SD":
                                        aturan_opsi = "3 opsi (A sampai C)"
                                    elif jenjang == "SMP":
                                        aturan_opsi = "4 opsi (A sampai D)"
                                    else:
                                        aturan_opsi = "5 opsi (A sampai E)"

                                    prompt = f"""Bertindaklah sebagai pakar kurikulum dan penyusun instrumen asesmen profesional. Buatkan {jumlah_soal} butir soal dengan bentuk **{jenis_soal}**, dalam bentuk asesmen {sub_asesmen} ({jenis_asesmen}) untuk Mata Pelajaran: {mapel}, Materi/Topik: {materi}, Jenjang: {jenjang} ({fase}, {kelas}). 
                                    
                                    Ketentuan Khusus:
                                    1. Gunakan pendekatan Pembelajaran Mendalam (Deep Learning) yang merangsang berpikir kritis dan kontekstual.
                                    2. Jika jenis soal adalah Pilihan Ganda, gunakan {aturan_opsi}.
                                    3. Berikan kunci jawaban yang jelas serta pembahasan mendalam untuk setiap soal."""

                                    model = genai.GenerativeModel("gemini-3.5-flash")
                                    response = model.generate_content(prompt)
                                    content_text = response.text

                                    docx_file = generate_professional_word_document(
                                        mapel=mapel,
                                        materi=materi,
                                        kelas=kelas,
                                        jenjang=jenjang,
                                        fase=fase,
                                        jenis_asesmen=jenis_asesmen,
                                        sub_asesmen=sub_asesmen,
                                        jumlah_soal=jumlah_soal,
                                        jenis_soal=jenis_soal,
                                        content_text=content_text,
                                    )

                                    st.session_state[f"file_bytes_{idx}"] = docx_file.getvalue()
                                    st.session_state[f"file_name_{idx}"] = f"Asesmen_{mapel}_{materi}.docx"
                                    st.success("✅ Dokumen siap diunduh!")
                                except Exception as e:
                                    st.error(f"Gagal membuat dokumen: {e}")

                        # Jika file sudah di-generate, tampilkan tombol unduh file yang sesungguhnya
                        if f"file_bytes_{idx}" in st.session_state:
                            st.download_button(
                                label="💾 Simpan File Word",
                                data=st.session_state[f"file_bytes_{idx}"],
                                file_name=st.session_state[f"file_name_{idx}"],
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                key=f"dl_btn_ready_{idx}",
                                use_container_width=True
                            )
                    st.markdown("---")

# --- MENU 4: INPUT DAN REKAP NILAI SISWA ---
elif menu == "📊 Input dan Rekap Nilai Siswa":
    st.subheader("Simpan & Rekap Nilai Hasil Asesmen")

    if not user_spreadsheet_id:
        st.error("⚠️ Spreadsheet ID untuk akun Anda tidak ditemukan di Master Registry. Hubungi administrator.")
        st.stop()

    @st.cache_data(ttl=10)
    def fetch_master_data(sheet_id):
        try:
            ss = gc.open_by_key(sheet_id)
            ws = ss.worksheet("Data Kelas-Siswa")
            records = ws.get_all_records()
            return records
        except Exception:
            return []

    master_data = fetch_master_data(user_spreadsheet_id)

    daftar_sekolah = []
    for r in master_data:
        sekolah_val = str(r.get("Sekolah", "")).strip()
        if sekolah_val and sekolah_val not in daftar_sekolah:
            daftar_sekolah.append(sekolah_val)

    if not daftar_sekolah:
        daftar_sekolah = ["SMK Negeri 2 Bangkalan"]

    if "val_mapel" not in st.session_state:
        st.session_state.val_mapel = ""
    if "val_sekolah" not in st.session_state:
        st.session_state.val_sekolah = daftar_sekolah[0]
    if "val_kelas" not in st.session_state:
        st.session_state.val_kelas = ""
    if "val_jenis" not in st.session_state:
        st.session_state.val_jenis = "Asesmen Formatif"
    if "val_materi" not in st.session_state:
        st.session_state.val_materi = ""

    col_r1, col_r2 = st.columns(2)

    with col_r1:
        tanggal_input = st.date_input("Pilih Tanggal", value=datetime.now())
        r_mapel = st.text_input("Mata Pelajaran", value=st.session_state.val_mapel, placeholder="Contoh: Bahasa Indonesia")
        r_sekolah = st.selectbox("Pilih Sekolah", daftar_sekolah, key="input_sekolah")

        kelas_filtered = [r for r in master_data if str(r.get("Sekolah", "")).strip() == r_sekolah]
        daftar_kelas = []
        for r in kelas_filtered:
            kelas_val = str(r.get("Kelas", "")).strip()
            if kelas_val and kelas_val not in daftar_kelas:
                daftar_kelas.append(kelas_val)

        if not daftar_kelas:
            daftar_kelas = ["X TKR-1", "X DKV-1"]

        r_kelas = st.selectbox("Pilih Kelas", sorted(daftar_kelas), key="input_kelas")
        r_jenis = st.selectbox("Jenis Asesmen", ["Asesmen Formatif", "Asesmen Sumatif"], key="input_jenis")

    with col_r2:
        r_materi = st.text_input("Materi / Topik", value=st.session_state.val_materi, placeholder="Contoh: Teks LHO")

        siswa_filtered = [
            r for r in master_data 
            if str(r.get("Sekolah", "")).strip() == r_sekolah and str(r.get("Kelas", "")).strip() == r_kelas
        ]

        mapping_absen_nama = {}
        for r in siswa_filtered:
            try:
                absen_raw = r.get("No_Absen", r.get("No Absen", 1))
                nama_raw = r.get("Nama_Siswa", r.get("Nama Siswa", r.get("Nama", "")))
                absen = int(absen_raw)
                nama = str(nama_raw).strip()
                if nama:
                    mapping_absen_nama[absen] = nama
            except:
                continue

        if not mapping_absen_nama:
            mapping_absen_nama = {1: "Siswa 1", 2: "Siswa 2"}

        list_absen = sorted(list(mapping_absen_nama.keys()))
        r_no_absen = st.selectbox("Pilih No. Absen Siswa", list_absen)

        r_nama_siswa = mapping_absen_nama.get(r_no_absen, "")
        st.text_input("Nama Siswa (Otomatis dari Spreadsheet)", value=r_nama_siswa, disabled=True)

        r_nilai = st.number_input("Nilai Siswa", min_value=0, max_value=100, value=80)

        if r_jenis == "Asesmen Formatif":
            sub_asesmen_input_options = ["Formatif Tertulis", "Formatif Tidak Tertulis"]
        else:
            sub_asesmen_input_options = ["Tulis", "Lisan", "Tugas", "Praktik", "Proyek", "Produk"]

        r_sub_jenis = st.selectbox("Bentuk / Sub Jenis Asesmen", sub_asesmen_input_options, key="input_sub_jenis")

    if st.button("💾 Simpan Nilai ke Google Sheets", use_container_width=True):
        if not r_mapel or not r_nama_siswa:
            st.warning("Mohon lengkapi Mata Pelajaran dan pastikan Nama Siswa terpilih.")
        else:
            try:
                ss = gc.open_by_key(user_spreadsheet_id)
                correct_header = [
                    "Tanggal", "Sekolah", "Mata Pelajaran", "Kelas", "Jenis Asesmen", 
                    "Sub Jenis Asesmen", "Materi", "No Absen", "Nama Siswa", "Nilai"
                ]

                try:
                    sheet = ss.worksheet("Rekap_Nilai")
                    sheet.update("A1:J1", [correct_header])
                except:
                    sheet = ss.add_worksheet(title="Rekap_Nilai", rows=100, cols=10)
                    sheet.append_row(correct_header)

                tanggal_str = tanggal_input.strftime("%Y-%m-%d")
                sheet.append_row([
                    tanggal_str, r_sekolah, r_mapel, r_kelas, r_jenis, 
                    r_sub_jenis, r_materi, r_no_absen, r_nama_siswa, r_nilai
                ])

                st.session_state.val_mapel = r_mapel
                st.session_state.val_sekolah = r_sekolah
                st.session_state.val_kelas = r_kelas
                st.session_state.val_jenis = r_jenis
                st.session_state.val_materi = r_materi

                st.success(f"🎉 Nilai untuk **{r_nama_siswa}** (Absen: {r_no_absen}) berhasil disimpan ke spreadsheet pribadi Anda!")
                st.balloons()
            except Exception as e:
                st.error(f"Gagal menyimpan ke database: {e}")

    # --- BAGIAN MENU UNDUH / EXPORT REKAP NILAI ---
    st.markdown("---")
    st.subheader("📥 Unduh Rekap Nilai Siswa")
    st.markdown("Pilih filter sekolah, kelas, dan mata pelajaran untuk mengunduh rekap nilai ke dalam format Excel.")

    try:
        ss_rekap = gc.open_by_key(user_spreadsheet_id)
        ws_rekap = ss_rekap.worksheet("Rekap_Nilai")
        all_rekap_rows = ws_rekap.get_all_values()
        if len(all_rekap_rows) > 1:
            df_rekap = pd.DataFrame(all_rekap_rows[1:], columns=all_rekap_rows[0])
        else:
            df_rekap = pd.DataFrame()
    except:
        df_rekap = pd.DataFrame()

    if not df_rekap.empty:
        list_dl_sekolah = df_rekap["Sekolah"].unique().tolist() if "Sekolah" in df_rekap.columns else daftar_sekolah
        dl_sekolah = st.selectbox("Filter Sekolah untuk Unduh", list_dl_sekolah, key="dl_sek")

        df_filtered_sek = df_rekap[df_rekap["Sekolah"] == dl_sekolah] if "Sekolah" in df_rekap.columns else df_rekap
        list_dl_kelas = df_filtered_sek["Kelas"].unique().tolist() if "Kelas" in df_filtered_sek.columns else []
        dl_kelas = st.selectbox("Filter Kelas untuk Unduh", list_dl_kelas if list_dl_kelas else ["Semua Kelas"], key="dl_kls")

        df_filtered_kls = df_filtered_sek[df_filtered_sek["Kelas"] == dl_kelas] if dl_kelas != "Semua Kelas" else df_filtered_sek
        list_dl_mapel = df_filtered_kls["Mata Pelajaran"].unique().tolist() if "Mata Pelajaran" in df_filtered_kls.columns else []
        dl_mapel = st.selectbox("Filter Mata Pelajaran untuk Unduh", list_dl_mapel if list_dl_mapel else ["Semua Mapel"], key="dl_mpl")

        df_final_dl = df_filtered_kls
        if dl_mapel != "Semua Mapel":
            df_final_dl = df_final_dl[df_final_dl["Mata Pelajaran"] == dl_mapel]

        st.write(f"📊 Menemukan **{len(df_final_dl)}** data nilai sesuai filter yang dipilih.")
        if not df_final_dl.empty:
            st.dataframe(df_final_dl, use_container_width=True)

            output_excel = BytesIO()
            with pd.ExcelWriter(output_excel, engine="openpyxl") as writer:
                df_final_dl.to_excel(writer, index=False, sheet_name="Rekap_Nilai")
            output_excel.seek(0)

            st.download_button(
                label="📥 Download Data Rekap Terpilih (.xlsx)",
                data=output_excel,
                file_name=f"Rek_Nilai_{dl_sekolah}_{dl_kelas}_{dl_mapel}.xlsx".replace(" ", "_"),
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                use_container_width=True,
            )
        else:
            st.info("Tidak ada data nilai yang cocok dengan kombinasi filter tersebut.")
    else:
        st.info("Belum ada data rekap nilai yang tersimpan di spreadsheet Anda. Silakan simpan beberapa nilai terlebih dahulu.")
