from datetime import datetime
from io import BytesIO
import json
import docx
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.shared import Inches, Pt, RGBColor
import google.generativeai as genai
import streamlit as st
from styles import apply_global_styles

# Konfigurasi Halaman[cite: 5]
st.set_page_config(
    page_title="Otomatisasi Penyusunan Modul Ajar PM",
    page_icon="📚",
    layout="wide",
)
apply_global_styles()

with st.sidebar:
  st.markdown(
      f"""
        <div class="user-profile-box">
            <span style="font-size: 24px;">👨‍💻</span><br>
            <b style="color: #facc15; font-size: 14px;">{st.session_state.get('guru_nama', 'Guru')}</b><br>
            <span style="color: #94a3b8; font-size: 11px;">Sesi Aktif & Terverifikasi</span>
        </div>
        """,
      unsafe_allow_html=True,
  )

st.markdown(
    """
    <style>
        .block-container {
            padding-top: 1.5rem;
            padding-bottom: 1rem;
        }
    </style>
""",
    unsafe_allow_html=True,
)

if "logged_in" not in st.session_state or not st.session_state.logged_in:
  st.warning(
      "⚠️ Anda belum login. Silakan kembali ke Halaman Utama (`app.py`) untuk"
      " masuk ke portal terlebih dahulu."
  )
  st.stop()

try:
  api_key_default = st.secrets.get("GEMINI_API_KEY", "")
except Exception:
  api_key_default = ""

st.markdown(
    """
    <style>
    .stApp {
        background-color: #0e1117;
        font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
    }
    .header-card {
        background: linear-gradient(135deg, #1e293b 0%, #0f172a 100%);
        padding: 20px 25px;
        border-radius: 12px;
        border: 1px solid #334155;
        margin-bottom: 20px;
        box-shadow: 0 10px 25px -5px rgba(0, 0, 0, 0.3);
    }
    .header-title {
        color: #f8fafc;
        font-size: 22px;
        font-weight: 700;
        margin: 0;
        letter-spacing: 0.3px;
        text-align: center;
    }
    @keyframes blink-animation {
        0% { opacity: 1; color: #facc15; }
        50% { opacity: 0.35; color: #38bdf8; }
        100% { opacity: 1; color: #facc15; }
    }
    .header-subtitle {
        font-size: 13px;
        margin-top: 8px;
        margin-bottom: 0;
        text-align: center;
        animation: blink-animation 1.6s infinite ease-in-out;
        font-weight: 600;
    }
    .section-header {
        font-size: 1.2rem;
        font-weight: 700;
        color: #f8fafc;
        margin-top: 1.2rem;
        margin-bottom: 0.8rem;
        white-space: nowrap;
        letter-spacing: 0.2px;
    }
    .stButton > button {
        width: 100%;
        border-radius: 8px;
        font-weight: 600;
        background: linear-gradient(135deg, #4f46e5 0%, #3b82f6 100%);
        color: white;
        border: none;
        padding: 0.75rem 1rem;
        box-shadow: 0 4px 12px rgba(59, 130, 246, 0.35);
        transition: all 0.3s ease;
    }
    .stButton > button:hover {
        background: linear-gradient(135deg, #4338ca 0%, #2563eb 100%);
        box-shadow: 0 6px 18px rgba(59, 130, 246, 0.5);
        transform: translateY(-2px);
    }
    </style>
""",
    unsafe_allow_html=True,
)


def set_cell_background(cell, fill_color):
  shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_color}"/>')
  cell._tc.get_or_add_tcPr().append(shading_elm)


def generate_docx(
    data_ai,
    nama_sekolah,
    semester,
    tahun_pelajaran,
    mata_pelajaran,
    fase_kelas,
    topik,
    alokasi_waktu,
    pertemuan_ke,
    nama_penulis,
    nama_kota,
    tanggal_pembuatan,
    nip_penulis,
):
  doc = docx.Document()

  for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

  style = doc.styles["Normal"]
  font = style.font
  font.name = "Arial"
  font.size = Pt(10)
  font.color.rgb = RGBColor(51, 51, 51)

  p_title = doc.add_paragraph()
  p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
  p_title.paragraph_format.space_before = Pt(0)
  p_title.paragraph_format.space_after = Pt(12)
  run_title = p_title.add_run("MODUL AJAR PEMBELAJARAN MENDALAM")
  run_title.font.name = "Arial"
  run_title.font.size = Pt(15)
  run_title.font.bold = True
  run_title.font.color.rgb = RGBColor(74, 46, 33)

  def add_section_table(title_text, rows_data):
    table = doc.add_table(rows=len(rows_data) + 1, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    hdr_cells = table.rows[0].cells
    hdr_cells[0].merge(hdr_cells[1])
    hdr_cells[0].text = title_text
    set_cell_background(hdr_cells[0], "5A3825")
    for p in hdr_cells[0].paragraphs:
      p.alignment = WD_ALIGN_PARAGRAPH.LEFT
      p.paragraph_format.space_before = Pt(4)
      p.paragraph_format.space_after = Pt(4)
      for run in p.runs:
        run.font.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(255, 255, 255)

    for idx, (label, val) in enumerate(rows_data):
      row_cells = table.rows[idx + 1].cells
      row_cells[0].text = label
      row_cells[0].width = Inches(2.3)
      row_cells[1].width = Inches(4.2)
      set_cell_background(row_cells[0], "F5EBE0")

      for p in row_cells[0].paragraphs:
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.15
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in p.runs:
          run.font.size = Pt(10)
          run.font.bold = True
          run.font.color.rgb = RGBColor(51, 51, 51)

      val_str = str(val).replace("LKPD", "LKM").replace(
          "Lembar Kegiatan Murid", "Lembar Kerja Murid"
      )
      row_cells[1].text = ""

      lines = val_str.split("\n")
      for line_idx, line in enumerate(lines):
        if line_idx == 0:
          p_right = row_cells[1].paragraphs[0]
        else:
          p_right = row_cells[1].add_paragraph()

        p_right.paragraph_format.space_before = Pt(4)
        p_right.paragraph_format.space_after = Pt(4)
        p_right.paragraph_format.line_spacing = 1.15
        p_right.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        if ":" in line:
          parts = line.split(":", 1)
          prefix = parts[0].strip() + ":"
          content = parts[1].strip()

          r_prefix = p_right.add_run(prefix + " ")
          r_prefix.font.size = Pt(10)
          r_prefix.font.bold = True
          r_prefix.font.color.rgb = RGBColor(51, 51, 51)

          r_content = p_right.add_run(content)
          r_content.font.size = Pt(10)
          r_content.font.bold = False
          r_content.font.color.rgb = RGBColor(51, 51, 51)
        else:
          r_normal = p_right.add_run(line)
          r_normal.font.size = Pt(10)
          r_normal.font.bold = False
          r_normal.font.color.rgb = RGBColor(51, 51, 51)

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

  tabel_identifikasi = [
      ("Penulis Modul", nama_penulis),
      ("Satuan Pendidikan", nama_sekolah),
      ("Mata Pelajaran", mata_pelajaran),
      ("Fase / Kelas", fase_kelas),
      ("Semester / Tahun Pelajaran", f"{semester} / {tahun_pelajaran}"),
      ("Materi / Topik", topik),
      ("Alokasi Waktu", alokasi_waktu),
      ("Pertemuan Ke-", pertemuan_ke),
  ]
  add_section_table("IDENTIFIKASI DAN INFORMASI UMUM", tabel_identifikasi)

  tabel_dpl = [
      (
          "Dimensi Profil Lulusan",
          data_ai.get(
              "dimensi_profil_lulusan",
              "☑ Penalaran Kritis: Peserta didik dilatih menganalisis masalah"
              " secara logis.\n☑ Kolaborasi: Bekerja sama dalam kelompok"
              " investigasi.\n☑ Kemandirian: Bertanggung jawab atas tugas"
              " mandiri.\n☑ Komunikasi: Mempresentasikan hasil kerja.",
          ),
      ),
  ]
  add_section_table("DIMENSI PROFIL LULUSAN", tabel_dpl)

  tabel_tujuan = [
      (
          "Tujuan Pembelajaran",
          data_ai.get(
              "tujuan_pembelajaran",
              "Peserta didik mampu menguasai kompetensi sesuai materi.",
          ),
      ),
  ]
  add_section_table("TUJUAN PEMBELAJARAN", tabel_tujuan)

  tabel_pemahaman = [
      (
          "Pemahaman Bermakna",
          data_ai.get(
              "pemahaman_bermakna",
              "Manfaat praktis dan esensi pembelajaran bagi kehidupan.",
          ),
      ),
      (
          "Pertanyaan Pemantik",
          data_ai.get(
              "pertanyaan_pemantik",
              "Pertanyaan kritis untuk menstimulasi rasa ingin tahu peserta"
              " didik.",
          ),
      ),
  ]
  add_section_table(
      "PEMAHAMAN BERMAKNA & PERTANYAAN PEMANTIK", tabel_pemahaman
  )

  tabel_kerangka = [
      (
          "Praktik Pedagogis",
          data_ai.get(
              "praktik_pedagogis",
              "Model Pembelajaran: Problem Based Learning\nMetode"
              " Pembelajaran Pendukung: Diskusi, Tanya Jawab, Analisis Teks",
          ),
      ),
      (
          "Kemitraan Pembelajaran",
          data_ai.get(
              "kemitraan_pembelajaran",
              "Kemitraan Lingkungan Sekolah: Kolaborasi guru mapel"
              " produktif.\nKemitraan Lingkungan Luar Sekolah: Pemanfaatan"
              " data/narasumber instansi terkait.",
          ),
      ),
      (
          "Lingkungan Belajar",
          data_ai.get(
              "lingkungan_belajar",
              "Ruang Fisik: Kelas fleksibel dan kolaboratif.\nRuang Virtual:"
              " Google Drive / LMS Sekolah.\nBudaya Belajar: Kolaboratif,"
              " Berpikir Kritis, Keterbukaan.",
          ),
      ),
      (
          "Pemanfaatan Digital",
          data_ai.get(
              "pemanfaatan_digital",
              "Tahap Perencanaan: AI & Cloud Storage.\nTahap Pelaksanaan: QR"
              " Code & Audio/Video Digital.\nTahap Asesmen: Google Form /"
              " Menti.",
          ),
      ),
  ]
  add_section_table("KERANGKA PEMBELAJARAN", tabel_kerangka)

  tabel_pengalaman = [
      (
          "Kegiatan Pendahuluan",
          data_ai.get(
              "kegiatan_pendahuluan",
              "Orientasi, Apersepsi, Motivasi, dan Asesmen Diagnostik awal.",
          ),
      ),
      (
          "Kegiatan Inti (Memahami)",
          data_ai.get(
              "kegiatan_memahami",
              "Eksplorasi konsep dan penyajian masalah autentik.",
          ),
      ),
      (
          "Kegiatan Inti (Mengaplikasi)",
          data_ai.get(
              "kegiatan_mengaplikasi",
              "Penyelidikan kolaboratif dan penerapan konsep dalam LKM.",
          ),
      ),
      (
          "Kegiatan Inti (Merefleksi)",
          data_ai.get(
              "kegiatan_merefleksi",
              "Presentasi kelompok, umpan balik konstruktif, dan penguatan.",
          ),
      ),
      (
          "Kegiatan Penutup",
          data_ai.get(
              "kegiatan_penutup",
              "Refleksi bersama yang menyenangkan (joyful) dan bermakna.",
          ),
      ),
  ]
  add_section_table("PENGALAMAN BELAJAR (LANGKAH-LANGKAH)", tabel_pengalaman)

  tabel_asesmen = [
      (
          "Asesmen Awal",
          data_ai.get(
              "asesmen_awal", "Cek kesiapan sebelum masuk topik pembelajaran."
          ),
      ),
      (
          "Asesmen Proses (Formatif)",
          data_ai.get(
              "asesmen_formatif",
              "Pemantauan partisipasi, keaktifan, dan pemahaman selama"
              " kegiatan.",
          ),
      ),
      (
          "Asesmen Akhir (Sumatif)",
          data_ai.get(
              "asesmen_sumatif",
              "Evaluasi hasil berbasis unjuk kerja atau refleksi kedalaman"
              " konsep.",
          ),
      ),
  ]
  add_section_table("ASESMEN PEMBELAJARAN", tabel_asesmen)

  p_sign = doc.add_paragraph()
  p_sign.alignment = WD_ALIGN_PARAGRAPH.RIGHT
  p_sign.paragraph_format.space_before = Pt(14)
  p_sign.paragraph_format.space_after = Pt(4)
  p_sign.add_run(f"{nama_kota}, {tanggal_pembuatan}\nPenyusun,\n\n\n")
  run_name = p_sign.add_run(f"{nama_penulis}")
  run_name.font.bold = True
  p_sign.add_run(f"\nNIP. {nip_penulis}")

  # HALAMAN 2: RUBRIK PENILAIAN
  doc.add_page_break()
  p_rubrik_title = doc.add_paragraph()
  p_rubrik_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
  p_rubrik_title.paragraph_format.space_after = Pt(12)
  r_rub_t = p_rubrik_title.add_run("RUBRIK PENILAIAN & PEDOMAN PENSKORAN")
  r_rub_t.font.name = "Arial"
  r_rub_t.font.size = Pt(15)
  r_rub_t.font.bold = True
  r_rub_t.font.color.rgb = RGBColor(74, 46, 33)

  table_id_rubrik = doc.add_table(rows=3, cols=2)
  table_id_rubrik.style = "Table Grid"
  table_id_rubrik.alignment = WD_TABLE_ALIGNMENT.CENTER
  table_id_rubrik.rows[0].cells[0].text = "Nama Guru / Pengamat:"
  table_id_rubrik.rows[0].cells[1].text = f"{nama_penulis}"
  table_id_rubrik.rows[1].cells[0].text = "Kelas / Fase:"
  table_id_rubrik.rows[1].cells[1].text = f"{fase_kelas}"
  table_id_rubrik.rows[2].cells[0].text = "Mata Pelajaran / Topik:"
  table_id_rubrik.rows[2].cells[1].text = f"{mata_pelajaran} - {topik}"

  for row in table_id_rubrik.rows:
    row.cells[0].width = Inches(2.3)
    row.cells[1].width = Inches(4.2)
    set_cell_background(row.cells[0], "F5EBE0")
    for cell in row.cells:
      for p in cell.paragraphs:
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        for run in p.runs:
          run.font.size = Pt(10)
          run.font.bold = True

  doc.add_paragraph().paragraph_format.space_after = Pt(6)

  p_sub = doc.add_paragraph()
  run_sub = p_sub.add_run("A. Rubrik Penilaian Kinerja / Kompetensi")
  run_sub.font.bold = True
  run_sub.font.size = Pt(10.5)
  run_sub.font.color.rgb = RGBColor(74, 46, 33)

  rubrik_data = data_ai.get("rubrik_penilaian", {})
  if isinstance(rubrik_data, dict) and rubrik_data:
    rubrik_table = doc.add_table(rows=len(rubrik_data) + 1, cols=5)
    rubrik_table.style = "Table Grid"
    rubrik_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    hdr_cells = rubrik_table.rows[0].cells
    headers = [
        "Kriteria Penilaian",
        "Perlu Bimbingan",
        "Cukup",
        "Baik",
        "Sangat Baik",
    ]
    col_widths = [
        Inches(1.5),
        Inches(1.25),
        Inches(1.25),
        Inches(1.25),
        Inches(1.25),
    ]

    for idx, text_hdr in enumerate(headers):
      hdr_cells[idx].text = text_hdr
      hdr_cells[idx].width = col_widths[idx]
      set_cell_background(hdr_cells[idx], "5A3825")
      for p in hdr_cells[idx].paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        for run in p.runs:
          run.font.bold = True
          run.font.size = Pt(9.5)
          run.font.color.rgb = RGBColor(255, 255, 255)

    for row_idx, (k, v) in enumerate(rubrik_data.items()):
      row_cells = rubrik_table.rows[row_idx + 1].cells
      if isinstance(v, dict):
        nama = v.get("nama_kriteria", k)
        pb = v.get("perlu_bimbingan", "-")
        c = v.get("cukup", "-")
        b = v.get("baik", "-")
        sb = v.get("sangat_baik", "-")
      else:
        nama = str(k)
        pb, c, b, sb = str(v), "", "", ""

      row_values = [nama, pb, c, b, sb]
      for col_idx, val_text in enumerate(row_values):
        row_cells[col_idx].text = str(val_text)
        row_cells[col_idx].width = col_widths[col_idx]

        if col_idx == 0:
          set_cell_background(row_cells[col_idx], "F5EBE0")

        for p in row_cells[col_idx].paragraphs:
          p.paragraph_format.space_before = Pt(4)
          p.paragraph_format.space_after = Pt(4)
          p.paragraph_format.line_spacing = 1.15
          p.alignment = WD_ALIGN_PARAGRAPH.LEFT
          for run in p.runs:
            run.font.size = Pt(9.0)
            run.font.bold = (col_idx == 0)
            run.font.color.rgb = RGBColor(51, 51, 51)

  doc.add_paragraph().paragraph_format.space_after = Pt(6)

  # HALAMAN 3: INSTRUMEN ASESMEN PROSES (FORMATIF)
  doc.add_page_break()
  p_inst_title = doc.add_paragraph()
  p_inst_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
  p_inst_title.paragraph_format.space_after = Pt(12)
  r_inst_t = p_inst_title.add_run("INSTRUMEN ASESMEN PROSES (FORMATIF)")
  r_inst_t.font.name = "Arial"
  r_inst_t.font.size = Pt(15)
  r_inst_t.font.bold = True
  r_inst_t.font.color.rgb = RGBColor(74, 46, 33)

  table_id_inst = doc.add_table(rows=3, cols=2)
  table_id_inst.style = "Table Grid"
  table_id_inst.alignment = WD_TABLE_ALIGNMENT.CENTER
  table_id_inst.rows[0].cells[0].text = "Nama Guru / Pengamat:"
  table_id_inst.rows[0].cells[1].text = f"{nama_penulis}"
  table_id_inst.rows[1].cells[0].text = "Kelas / Fase:"
  table_id_inst.rows[1].cells[1].text = f"{fase_kelas}"
  table_id_inst.rows[2].cells[0].text = "Mata Pelajaran / Topik:"
  table_id_inst.rows[2].cells[1].text = f"{mata_pelajaran} - {topik}"

  for row in table_id_inst.rows:
    row.cells[0].width = Inches(2.3)
    row.cells[1].width = Inches(4.2)
    set_cell_background(row.cells[0], "F5EBE0")
    for cell in row.cells:
      for p in cell.paragraphs:
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        for run in p.runs:
          run.font.size = Pt(10)
          run.font.bold = True

  doc.add_paragraph().paragraph_format.space_after = Pt(6)
  instrumen_data = data_ai.get("instrumen_formatif", {})
  if isinstance(instrumen_data, dict) and instrumen_data:
    inst_rows = []
    for inst_k, inst_v in instrumen_data.items():
      label_text = inst_k.replace("_", " ").title()
      inst_rows.append((label_text, str(inst_v)))
    add_section_table("LEMBAR OBSERVASI / FORMATIF KELAS", inst_rows)

  # HALAMAN 4: BAHAN AJAR (BARU DITAMBAHKAN)
  doc.add_page_break()
  p_bahan_title = doc.add_paragraph()
  p_bahan_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
  p_bahan_title.paragraph_format.space_after = Pt(12)
  r_bahan_t = p_bahan_title.add_run("BAHAN AJAR / MATERI PEMBELAJARAN")
  r_bahan_t.font.name = "Arial"
  r_bahan_t.font.size = Pt(15)
  r_bahan_t.font.bold = True
  r_bahan_t.font.color.rgb = RGBColor(74, 46, 33)

  table_id_bahan = doc.add_table(rows=3, cols=2)
  table_id_bahan.style = "Table Grid"
  table_id_bahan.alignment = WD_TABLE_ALIGNMENT.CENTER
  table_id_bahan.rows[0].cells[0].text = "Mata Pelajaran:"
  table_id_bahan.rows[0].cells[1].text = f"{mata_pelajaran}"
  table_id_bahan.rows[1].cells[0].text = "Fase / Kelas / Topik:"
  table_id_bahan.rows[1].cells[1].text = f"{fase_kelas} - {topik}"
  table_id_bahan.rows[2].cells[0].text = "Alokasi Waktu / Pertemuan:"
  table_id_bahan.rows[2].cells[1].text = f"{alokasi_waktu} ({pertemuan_ke})"

  for row in table_id_bahan.rows:
    row.cells[0].width = Inches(2.3)
    row.cells[1].width = Inches(4.2)
    set_cell_background(row.cells[0], "F5EBE0")
    for cell in row.cells:
      for p in cell.paragraphs:
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        for run in p.runs:
          run.font.size = Pt(10)
          run.font.bold = True

  doc.add_paragraph().paragraph_format.space_after = Pt(6)
  bahan_data = data_ai.get("bahan_ajar", {})
  if isinstance(bahan_data, dict) and bahan_data:
    bahan_rows = []
    for b_k, b_v in bahan_data.items():
      label_text = b_k.replace("_", " ").title()
      bahan_rows.append((label_text, str(b_v)))
    add_section_table("URAIAN MATERI & KONSEP PEMBELAJARAN", bahan_rows)

  # HALAMAN 5: LEMBAR KERJA MURID (LKM)
  doc.add_page_break()
  p_lkm_title = doc.add_paragraph()
  p_lkm_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
  p_lkm_title.paragraph_format.space_after = Pt(12)
  r_lkm_t = p_lkm_title.add_run("LEMBAR KERJA MURID (LKM)")
  r_lkm_t.font.name = "Arial"
  r_lkm_t.font.size = Pt(15)
  r_lkm_t.font.bold = True
  r_lkm_t.font.color.rgb = RGBColor(74, 46, 33)

  table_id_lkm = doc.add_table(rows=3, cols=2)
  table_id_lkm.style = "Table Grid"
  table_id_lkm.alignment = WD_TABLE_ALIGNMENT.CENTER
  table_id_lkm.rows[0].cells[0].text = "Nama Kelompok / Peserta Didik:"
  table_id_lkm.rows[0].cells[1].text = "........................................"
  table_id_lkm.rows[1].cells[0].text = "Kelas / Fase:"
  table_id_lkm.rows[1].cells[1].text = f"{fase_kelas}"
  table_id_lkm.rows[2].cells[0].text = "Mata Pelajaran / Topik:"
  table_id_lkm.rows[2].cells[1].text = f"{mata_pelajaran} - {topik}"

  for row in table_id_lkm.rows:
    row.cells[0].width = Inches(2.3)
    row.cells[1].width = Inches(4.2)
    set_cell_background(row.cells[0], "F5EBE0")
    for cell in row.cells:
      for p in cell.paragraphs:
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        for run in p.runs:
          run.font.size = Pt(10)
          run.font.bold = True

  doc.add_paragraph().paragraph_format.space_after = Pt(6)
  lkm_data = data_ai.get("lkm_content", {})
  if isinstance(lkm_data, dict) and lkm_data:
    lkm_rows = []
    for lkm_k, lkm_v in lkm_data.items():
      label_text = lkm_k.replace("_", " ").title()
      lkm_rows.append((label_text, str(lkm_v)))
    add_section_table("STRUKTUR LEMBAR KERJA MURID (LKM)", lkm_rows)

  bio = BytesIO()
  doc.save(bio)
  bio.seek(0)
  return bio


st.markdown(
    """
    <div class="header-card">
        <h2 class="header-title">
            Otomatisasi Penyusunan Modul Ajar PM
        </h2>
        <div class="header-subtitle">
            <b>Pengembang:</b> Yustinus Budi Setyanta - PS Cabdin Bangkalan &nbsp;&nbsp;
            <em>Aplikasi Otomatisasi Perancangan Pembelajaran Mendalam</em>
        </div>
    </div>
    """,
    unsafe_allow_html=True,
)

st.markdown("---")

st.markdown(
    '<div class="section-header">⚙️ Parameter Pembelajaran</div>',
    unsafe_allow_html=True,
)
api_key = st.text_input(
    "Masukkan Google Gemini API Key", value="", type="password"
)  # <-- Baris 675-677

st.markdown(
    '💡 *Belum punya API Key? <a'
    ' href="https://aistudio.google.com/app/apikey" target="_blank">Klik di'
    ' sini untuk membuat secara mandiri & gratis</a>*',
    unsafe_allow_html=True,
)

col_param1, col_param2 = st.columns(2)

with col_param1:
  jenjang_pendidikan = st.selectbox(
      "Pilih Jenjang Pendidikan", ["SD / MI", "SMP / MTs", "SMA / MA", "SMK / MAK"]
  )

  if jenjang_pendidikan == "SD / MI":
    default_mapel = "Tematik / Kelas"
    jp_guidance = "Panduan: 1 JP = 35 Menit"
    fase_options = [
        "Fase A / Kelas 1 SD",
        "Fase A / Kelas 2 SD",
        "Fase B / Kelas 3 SD",
        "Fase B / Kelas 4 SD",
        "Fase C / Kelas 5 SD",
        "Fase C / Kelas 6 SD",
    ]
  elif jenjang_pendidikan == "SMP / MTs":
    default_mapel = "Matematika / IPA / IPS"
    jp_guidance = "Panduan: 1 JP = 40 Menit"
    fase_options = [
        "Fase D / Kelas 7 SMP",
        "Fase D / Kelas 8 SMP",
        "Fase D / Kelas 9 SMP",
    ]
  elif jenjang_pendidikan == "SMA / MA":
    default_mapel = "Bahasa Indonesia / Matematika"
    jp_guidance = "Panduan: 1 JP = 45 Menit"
    fase_options = [
        "Fase E / Kelas X SMA",
        "Fase F / Kelas XI SMA",
        "Fase F / Kelas XII SMA",
    ]
  else:
    default_mapel = (
        "Dasar-dasar Teknik Otomotif / Produk Kreatif dan Kewirausahaan"
    )
    jp_guidance = "Panduan: 1 JP = 45 Menit"
    fase_options = [
        "Fase E / Kelas X SMK (Program Dasar Keahlian)",
        "Fase F / Kelas XI SMK (Konsentrasi Keahlian)",
        "Fase F / Kelas XII SMK (Konsentrasi Keahlian)",
    ]

  mata_pelajaran = st.text_input(
      "Mata Pelajaran / Program Kejuruan", default_mapel
  )
  fase_kelas = st.selectbox("Fase / Kelas", fase_options)

with col_param2:
  topik = st.text_input(
      "Topik / Materi Pokok / Elemen",
      (
          "Contoh: Pemeliharaan Sistem Rem Kendaraan Ringan"
          if jenjang_pendidikan == "SMK / MAK"
          else "Contoh: Menyimak Teks Laporan Observasi Secara Kritis"
      ),
  )
  st.caption(jp_guidance)
  alokasi_waktu = st.text_input("Alokasi Waktu", "2 JP (2 x 45 Menit)")
  pertemuan_ke = st.text_input("Pertemuan Ke-", "1 (Pertemuan Pertama)")

st.markdown("---")

col_id1, col_id2 = st.columns(2)

with col_id1:
  st.markdown(
      '<div class="section-header">🏫 Identitas Satuan Pendidikan</div>',
      unsafe_allow_html=True,
  )
  nama_sekolah = st.text_input(
      "Nama Sekolah", st.session_state.get("sekolah", "SMK Negeri 2 Bangkalan")
  )
  semester = st.selectbox("Semester", ["Ganjil", "Genap"])
  tahun_pelajaran = st.text_input("Tahun Pelajaran", "2026/2027")

with col_id2:
  st.markdown(
      '<div class="section-header">✍️ Identitas Pengesahan Dokumen</div>',
      unsafe_allow_html=True,
  )
  nama_kota = st.text_input("Nama Kota", "Bangkalan")
  tanggal_pembuatan = st.text_input(
      "Tanggal / Bulan / Tahun", datetime.today().strftime("%d %B %Y")
  )
  nama_penulis = st.text_input(
      "Nama Penulis Modul",
      st.session_state.get("guru_nama", "Yustinus Budi Setyanta, S.Pd., M.Pd."),
  )
  nip_penulis = st.text_input("NIP Penulis", "196908302005011003")

st.markdown("---")

st.markdown("### 🚀 Generator Modul Ajar Pembelajaran Mendalam")
st.markdown(
    "Pastikan parameter dan identitas di atas sudah terisi dengan benar, lalu"
    " klik tombol di bawah untuk mulai menyusun dokumen secara otomatis"
    " menggunakan AI."
)

st.markdown("<br>", unsafe_allow_html=True)

if st.button("🚀 Buat Modul Ajar Pembelajaran Mendalam", use_container_width=True):
  if not api_key:
    st.error("Mohon masukkan Google Gemini API Key terlebih dahulu.")
  elif not topik:
    st.warning("Mohon isi topik pembelajaran.")
  else:
    with st.spinner(
        f"{nama_penulis} sedang menyusun Modul Ajar Pembelajaran Mendalam..."
    ):
      genai.configure(api_key=api_key)
      model = genai.GenerativeModel("gemini-3.5-flash")

      prompt = f"""
          Bertindaklah sebagai pakar kurikulum profesional. Buatkan konten Modul Ajar Berbasis Pembelajaran Mendalam (Deep Learning) yang **SANGAT LENGKAP, DETAIL, DAN KOMPREHENSIF** untuk:
          - Jenjang: {jenjang_pendidikan} ({fase_kelas})
          - Mata Pelajaran: {mata_pelajaran}
          - Topik / Materi Pokok: {topik}
          - Alokasi Waktu: {alokasi_waktu}
          - Pertemuan Ke-: {pertemuan_ke}

          Ketentuan Penting:
          1. Dimensi Profil Lulusan: Pilih 2 hingga 4 dimensi yang PALING RELEVAN dari 8 dimensi berikut (Keimanan dan Ketaqwaan terhadap Tuhan Yang Maha Esa, Kewargaan, Penalaran Kritis, Kreativitas, Kolaborasi, Kemandirian, Kesehatan, Komunikasi). **SANGAT PENTING: Tuliskan dan tampilkan HANYA dimensi yang dipilih saja (dengan tanda centang ☑ dan uraian penjelasannya). JANGAN SAMA SEKALI menyebutkan atau menuliskan daftar dimensi lain yang tidak dipilih/tidak digunakan.**
          2. Praktik Pedagogis: Gunakan format label persis berikut (dengan tanda titik dua):
             - Model Pembelajaran: [Uraian model seperti Problem Based Learning / Discovery Learning / dll]
             - Metode Pembelajaran Pendukung: [Uraian metode, misal 1. Studi Kasus Riil: ... 2. Demonstrasi Interaktif: ... dst]
          3. Kemitraan Pembelajaran: Gunakan format label persis berikut:
             - Kemitraan Lingkungan Sekolah: [...]
             - Kemitraan Lingkungan Luar Sekolah: [...]
          4. Lingkungan Belajar: Gunakan format label persis berikut:
             - Ruang Fisik: [...]
             - Ruang Virtual: [...]
             - Budaya Belajar: [...]
          5. Pemanfaatan Digital: Gunakan format label persis berikut:
             - Tahap Perencanaan: [...]
             - Tahap Pelaksanaan: [...]
             - Tahap Asesmen: [...]
          6. Pengalaman Belajar harus terstruktur mencakup Kegiatan Pendahuluan, Kegiatan Inti (Memahami, Mengaplikasi, Merefleksi), dan Kegiatan Penutup (refleksi joyful dan bermakna). Gunakan istilah **LKM (Lembar Kerja Murid)** (BUKAN LKPD atau Lembar Kegiatan Murid) di seluruh uraian.
          7. Asesmen Pembelajaran mencakup Asesmen Awal, Asesmen Proses (Formatif), dan Asesmen Akhir (Sumatif) beserta Rubrik Penilaian dan Pedoman Penskorannya.
          8. **Instrumen Asesmen Proses (Formatif)**: Sediakan instrumen asesmen proses/formatif yang mendalam pada kunci `instrumen_formatif` yang terstruktur dengan sub-bagian penting bertanda titik dua agar mudah disajikan dalam bentuk tabel rapi pada halaman khusus.
          9. **Bahan Ajar**: Sediakan materi pembelajaran/bahan bacaan yang mendalam dan komprehensif sesuai topik pada kunci `bahan_ajar` yang mencakup pengantar konsep, uraian materi inti, serta contoh kontekstual dengan format sub-bagian berlabel titik dua agar tersaji rapi sebagai halaman khusus Bahan Ajar sebelum LKM.
          10. **LKM (Lembar Kerja Murid)**: Sediakan konten LKM yang mendalam pada kunci `lkm_content` yang mencakup judul, tujuan, petunjuk kerja, serta langkah-langkah tugas/investigasi peserta didik pada halaman terpisah paling akhir.

          Berikan output HANYA dalam format JSON valid yang memuat kunci-kunci berikut:
          {{
            "dimensi_profil_lulusan": "Hanya tuliskan dimensi profil lulusan yang dipilih saja (gunakan tanda ☑) beserta uraian penerapannya. JANGAN menuliskan dimensi yang tidak dipilih.",
            "tujuan_pembelajaran": "Uraian tujuan pembelajaran yang spesifik, operasional, dan terukur sesuai materi.",
            "pemahaman_bermakna": "Uraian pemahaman bermakna yang mendalam terkait materi.",
            "pertanyaan_pemantik": "2 pertanyaan pemantik yang kontekstual dan menantang daya nalar kritis siswa.",
            "praktik_pedagogis": "Model Pembelajaran: [Isi model]\\nMetode Pembelajaran Pendukung: [Isi metode dengan penomoran]",
            "kemitraan_pembelajaran": "Kemitraan Lingkungan Sekolah: [Isi]\\nKemitraan Lingkungan Luar Sekolah: [Isi]",
            "lingkungan_belajar": "Ruang Fisik: [Isi]\\nRuang Virtual: [Isi]\\nBudaya Belajar: [Isi]",
            "pemanfaatan_digital": "Tahap Perencanaan: [Isi]\\nTahap Pelaksanaan: [Isi]\\nTahap Asesmen: [Isi]",
            "kegiatan_pendahuluan": "Langkah rinci kegiatan pendahuluan (orientasi, apersepsi, asesmen awal).",
            "kegiatan_memahami": "Langkah rinci kegiatan inti pada tahap Memahami.",
            "kegiatan_mengaplikasi": "Langkah rinci kegiatan inti pada tahap Mengaplikasi menggunakan LKM.",
            "kegiatan_merefleksi": "Langkah rinci kegiatan inti pada tahap Merefleksi dan presentasi.",
            "kegiatan_penutup": "Langkah rinci kegiatan penutup yang joyful dan bermakna.",
            "asesmen_awal": "Uraian asesmen awal untuk cek kesiapan belajar.",
            "asesmen_formatif": "Uraian asesmen proses/formatif pemantauan partisipasi.",
            "asesmen_sumatif": "Uraian asesmen akhir/sumatif evaluasi unjuk kerja.",
            "rubrik_penilaian": {{
              "kriteria_1": {{
                "nama_kriteria": "Nama kriteria pertama sesuai kompetensi materi",
                "perlu_bimbingan": "Deskripsi tingkat perlu bimbingan",
                "cukup": "Deskripsi tingkat cukup",
                "baik": "Deskripsi tingkat baik",
                "sangat_baik": "Deskripsi tingkat sangat baik"
              }},
              "kriteria_2": {{
                "nama_kriteria": "Nama kriteria kedua",
                "perlu_bimbingan": "Deskripsi...",
                "cukup": "Deskripsi...",
                "baik": "Deskripsi...",
                "sangat_baik": "Deskripsi..."
              }}
            }},
            "pedoman_penskoran": {{
              "rumus_nilai": "Rumus perhitungan nilai akhir",
              "kategori_predikat": "Interval nilai dan predikat kelulusan"
            }},
            "instrumen_formatif": {{
              "judul_instrumen": "Judul spesifik instrumen asesmen proses",
              "tujuan_asesmen": "Tujuan penggunaan lembar asesmen formatif",
              "aspek_yang_diamati": "Indikator atau aspek keaktifan/proses yang dinilai dengan format sub-bagian berlabel titik dua",
              "pedoman_pengamatan": "Petunjuk penilaian atau rubrik ceklis observasi singkat"
            }},
            "bahan_ajar": {{
              "pengantar_konsep": "Definisi atau pengantar esensial mengenai topik pembelajaran",
              "uraian_materi_inti": "Penjelasan detail dan komprehensif substansi materi yang dipelajari",
              "contoh_kontekstual": "Studi kasus atau contoh nyata pengaplikasian materi dalam kehidupan sehari-hari"
            }},
            "lkm_content": {{
              "judul_LKM": "Judul spesifik LKM",
              "tujuan_LKM": "Tujuan pengerjaan LKM bagi peserta didik",
              "petunjuk_kerja": "Langkah panduan keselamatan dan cara pengerjaan dengan format sub-bagian berlabel titik dua",
              "tugas_analisis": "Rincian tugas investigasi, pertanyaan kerja, atau tabel isian praktik"
            }}
          }}
          """

      response = model.generate_content(prompt)
      text_resp = response.text.strip()

      if text_resp.startswith("```json"):
        text_resp = text_resp[7:]
      if text_resp.startswith("```"):
        text_resp = text_resp[3:]
      if text_resp.endswith("```"):
        text_resp = text_resp[:-3]
      text_resp = text_resp.strip()

      try:
        data_ai = json.loads(text_resp)
      except Exception:
        data_ai = {}

      st.success("🎉 Modul Ajar Sesuai Sistematika Baru Berhasil Disusun!")
      st.info(
          "Dokumen Word (.docx) siap diunduh lengkap dengan halaman terpisah"
          " untuk Rubrik & Pedoman, Instrumen Format, **Bahan Ajar**, serta"
          " Lembar Kerja Murid (LKM)."
      )

      docx_file = generate_docx(
          data_ai,
          nama_sekolah,
          semester,
          tahun_pelajaran,
          mata_pelajaran,
          fase_kelas,
          topik,
          alokasi_waktu,
          pertemuan_ke,
          nama_penulis,
          nama_kota,
          tanggal_pembuatan,
          nip_penulis,
      )

      st.download_button(
          label="📥 Unduh Modul Ajar Pembelajaran Mendalam (.docx)",
          data=docx_file,
          file_name=f"Modul_Ajar_{topik.replace(' ', '_')}.docx",
          mime=(
              "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
          ),
          use_container_width=True,
      )
